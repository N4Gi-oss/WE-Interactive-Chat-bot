"""
Loader for user-uploaded documents (Requirement: PDF, DOCX, TXT, HTML, Images).

Extracts raw text from each supported file type, then reuses the exact
same chunk_page() logic from scraper.py so uploaded content is chunked
consistently with the scraped website content.

Images: instead of local OCR (Tesseract, which needs a system-level
install), we send the image directly to Gemini's multimodal API and ask
it to transcribe/describe the visible text. No extra system dependency.

Uploaded chunks are tagged source_type="user_upload" (vs. "website" for
scraped pages), so the RAG layer can always tell where an answer came
from and, if needed, restrict search to one source type.
"""

import os
from pathlib import Path

from bs4 import BeautifulSoup

from scraper import chunk_page

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".html", ".htm", ".png", ".jpg", ".jpeg"}


class UnsupportedFileTypeError(Exception):
    pass


class TextExtractionError(Exception):
    """Raised when a file is a supported type but text could not be
    extracted (corrupted file, empty file, missing optional dependency,
    etc.)."""
    pass


def _extract_pdf(path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(path: str) -> str:
    import docx
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


def _extract_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _extract_html(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    texts = [
        tag.get_text(strip=True, separator=" ")
        for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "td"])
    ]
    return "\n".join(t for t in texts if len(t) > 5)


def _extract_image(path: str) -> str:
    """
    Uses Gemini's native multimodal input instead of local OCR: no
    Tesseract or system dependency needed, works the same on any OS.
    """
    from rag import get_gemini_client

    client = get_gemini_client()
    with open(path, "rb") as f:
        image_bytes = f.read()

    ext = Path(path).suffix.lower().lstrip(".")
    mime_type = "image/jpeg" if ext in ("jpg", "jpeg") else "image/png"

    response = client.models.generate_content(
        model="gemini-flash-lite-latest",
        contents=[
            {
                "role": "user",
                "parts": [
                    {"inline_data": {"mime_type": mime_type, "data": image_bytes}},
                    {"text": "استخرج كل النص المكتوب في الصورة دي حرفيًا (عربي أو إنجليزي). لو مفيش نص، اوصف محتوى الصورة باختصار."},
                ],
            }
        ],
    )
    return response.text or ""


_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".txt": _extract_txt,
    ".html": _extract_html,
    ".htm": _extract_html,
    ".png": _extract_image,
    ".jpg": _extract_image,
    ".jpeg": _extract_image,
}


def extract_text(file_path: str) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = Path(file_path).suffix.lower()
    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        raise UnsupportedFileTypeError(
            f"'{ext}' is not supported. Supported types: "
            f"{', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )
    text = extractor(file_path)
    if not text or not text.strip():
        raise TextExtractionError(
            f"No readable text could be extracted from '{os.path.basename(file_path)}'."
        )
    return text


def process_uploaded_file(file_path: str, category: str = "user_upload") -> list[dict]:
    """
    Full pipeline for one uploaded file: extract -> chunk -> build chunk
    dicts in the exact shape vector_store.add_chunks() expects.
    """
    filename = os.path.basename(file_path)
    text = extract_text(file_path)
    chunks = chunk_page(text)

    return [
        {
            "chunk_id": f"upload_{filename}_{i}",
            "url": filename,
            "text": chunk_text,
            "category": category,
            "source_type": "user_upload",
        }
        for i, chunk_text in enumerate(chunks)
    ]